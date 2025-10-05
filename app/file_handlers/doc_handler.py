"""
DOC/DOCX Handler: Extract text from Microsoft Word documents.
"""

from pathlib import Path
from typing import Dict, Any
from docx import Document


class DocHandler:
    """Handler for processing DOC and DOCX files."""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.doc']
    
    def extract(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a Word document.
        
        Args:
            file_path: Path to the DOC/DOCX file
            
        Returns:
            Dictionary containing text content and metadata
        """
        doc_path = Path(file_path)
        if not doc_path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")
        
        results = {
            'text_content': [],
            'images': [],
            'metadata': {
                'file_name': doc_path.name,
                'file_path': str(doc_path),
                'file_type': 'docx'
            }
        }
        
        try:
            # Open document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    full_text.append(table_text)
            
            # Combine all text
            if full_text:
                results['text_content'].append({
                    'text': '\n'.join(full_text)
                })
            
            # Update metadata
            results['metadata']['num_paragraphs'] = len(doc.paragraphs)
            results['metadata']['num_tables'] = len(doc.tables)
            
            # Extract core properties if available
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                results['metadata']['author'] = doc.core_properties.author
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                results['metadata']['title'] = doc.core_properties.title
            
        except Exception as e:
            raise Exception(f"Error processing document {file_path}: {str(e)}")
        
        return results
    
    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table in a Word document.
        
        Args:
            table: python-docx table object
            
        Returns:
            Formatted text from the table
        """
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_text.append(' | '.join(row_text))
        
        return '\n'.join(table_text)
    
    def is_supported(self, file_path: str) -> bool:
        """Check if the file extension is supported."""
        return Path(file_path).suffix.lower() in self.supported_extensions
